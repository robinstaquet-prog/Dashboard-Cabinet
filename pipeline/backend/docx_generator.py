"""
DOCX + ODT generator — produit les fichiers patient en mémoire.

Deux variantes sur disque (DOCX + ODT) :
  • save_normal()     → chiffré (AES-256 via crypto.py)
  • save_anonymized() → en clair (données pseudonymisées)

Sections dans l'ordre exact de la fiche papier.
"""
from __future__ import annotations

import io
import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Pt

from .crypto import encrypt_bytes

# LibreOffice pour conversion ODT
_SOFFICE = Path("C:/Program Files/LibreOffice/program/soffice.exe")


# ---------------------------------------------------------------------------
# Helpers internes
# ---------------------------------------------------------------------------

def _safe_name(prenom_nom: str) -> str:
    """'Marie Dupont' → 'Marie_Dupont'  (espaces → _, caractères valides conservés)."""
    s = prenom_nom.strip()
    s = re.sub(r"[^\w\s\-]", "", s)   # retire les caractères spéciaux sauf tiret
    return re.sub(r"\s+", "_", s)


def _v(value) -> str:
    """None → '—', sinon str(value)."""
    return "—" if value is None else str(value)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _field(doc: Document, label: str, value) -> None:
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(f"{label} : ")
    run.bold = True
    p.add_run(_v(value))


def _bold_cell(cell, text: str) -> None:
    """Met le texte en gras dans une cellule de tableau."""
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.bold = True


# ---------------------------------------------------------------------------
# Construction du document
# ---------------------------------------------------------------------------

def build_docx(patient: dict) -> bytes:
    """
    Génère un .docx en mémoire à partir du dict patient.
    Les champs null sont affichés comme '—'.
    Retourne les bytes du fichier .docx.
    """
    doc = Document()

    # ── Titre ────────────────────────────────────────────────────────────────
    identite = patient.get("identite") or {}
    name = identite.get("prenom_nom") or "Patient"
    doc.add_heading(name, level=0)

    _field(doc, "Date première séance", patient.get("date_premiere_seance"))
    doc.add_paragraph()

    # ── Identité ─────────────────────────────────────────────────────────────
    _heading(doc, "Identité")
    _field(doc, "Prénom / Nom",           identite.get("prenom_nom"))
    _field(doc, "Date de naissance",      identite.get("date_naissance"))
    _field(doc, "Profession",             identite.get("profession"))
    _field(doc, "Enfants",                identite.get("enfants"))
    _field(doc, "Téléphone",              identite.get("telephone"))
    _field(doc, "E-mail",                 identite.get("mail"))
    _field(doc, "Adresse",                identite.get("adresse"))
    _field(doc, "Canal de communication", identite.get("canal_communication"))

    # ── Motif de consultation ─────────────────────────────────────────────────
    _heading(doc, "Motif de consultation")
    doc.add_paragraph(_v(patient.get("motif_consultation")))

    # ── Symptôme ─────────────────────────────────────────────────────────────
    _heading(doc, "Symptôme")
    symptome = patient.get("symptome") or {}
    _field(doc, "Manifestation",              symptome.get("manifestation"))
    _field(doc, "Depuis",                     symptome.get("depuis"))
    _field(doc, "Aggravation / Amélioration", symptome.get("aggravation_amelioration"))

    # ── Histoire personnelle ──────────────────────────────────────────────────
    _heading(doc, "Histoire personnelle")
    doc.add_paragraph(_v(patient.get("histoire_personnelle")))

    # ── Antécédents ───────────────────────────────────────────────────────────
    _heading(doc, "Antécédents personnels et familiaux")
    doc.add_paragraph(_v(patient.get("antecedents_personnels_familiaux")))

    # ── Traitements ───────────────────────────────────────────────────────────
    _heading(doc, "Traitements en cours")
    doc.add_paragraph(_v(patient.get("traitements_en_cours")))

    # ── Stress / Fatigue ──────────────────────────────────────────────────────
    _heading(doc, "Stress / Fatigue psychique")
    doc.add_paragraph(_v(patient.get("stress_fatigue_psychique")))

    # ── Système digestif ──────────────────────────────────────────────────────
    _heading(doc, "Système digestif")
    doc.add_paragraph(_v(patient.get("systeme_digestif")))

    # ── Sommeil ───────────────────────────────────────────────────────────────
    _heading(doc, "Sommeil")
    doc.add_paragraph(_v(patient.get("sommeil")))

    # ── Cardio-vasculaire ─────────────────────────────────────────────────────
    _heading(doc, "Cardio-vasculaire")
    cardio = patient.get("cardio_vasculaire") or {}
    _field(doc, "Observations",   cardio.get("observations"))
    _field(doc, "Anticoagulants", cardio.get("anticoagulants"))

    # ── Cycle menstruel ───────────────────────────────────────────────────────
    _heading(doc, "Cycle menstruel")
    doc.add_paragraph(_v(patient.get("cycle_menstruel")))

    # ── Pouls / Langue ────────────────────────────────────────────────────────
    _heading(doc, "Pouls / Langue")
    doc.add_paragraph(_v(patient.get("pouls_langue")))

    # ── Notes découvertes en cours de suivi ──────────────────────────────────
    notes = patient.get("notes_en_cours_de_suivi") or []
    if notes:
        _heading(doc, "Notes découvertes en cours de suivi")
        for note in notes:
            doc.add_paragraph(f"• {_v(note)}", style="Normal")

    # ── Séances ───────────────────────────────────────────────────────────────
    _heading(doc, "Séances")
    seances = [s for s in (patient.get("seances") or [])
               if s.get("date") or s.get("remarques") or s.get("bilan_ttt")]

    if seances:
        for s in seances:
            num = s.get("numero", "")
            date = _v(s.get("date"))
            _heading(doc, f"Séance {num} — {date}", level=2)

            remarques = s.get("remarques") or s.get("bilan_ttt")
            if remarques:
                _field(doc, "Remarques", remarques)
            if s.get("pouls_langue"):
                _field(doc, "Pouls / Langue", s.get("pouls_langue"))
            if s.get("strategie"):
                _field(doc, "Stratégie", s.get("strategie"))

            points = s.get("points_utilises") or []
            if points:
                pts_str_parts = []
                for p in points:
                    if isinstance(p, str):
                        pts_str_parts.append(p)
                    else:
                        code = p.get("code", "")
                        tech = p.get("technique")
                        pts_str_parts.append(f"{code} {tech}" if tech else code)
                _field(doc, "Points", "  ·  ".join(pts_str_parts))

            if s.get("amelioration") is not None:
                _field(doc, "Amélioration", f"{s.get('amelioration')}/10")
            if s.get("a_faire_prochaine_seance"):
                _field(doc, "À faire prochaine séance", s.get("a_faire_prochaine_seance"))
            doc.add_paragraph()
    else:
        doc.add_paragraph("Aucune séance enregistrée.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Conversion ODT via LibreOffice
# ---------------------------------------------------------------------------

def _docx_to_odt_bytes(docx_bytes: bytes) -> bytes | None:
    """
    Convertit des bytes DOCX en bytes ODT via LibreOffice headless.
    Retourne None si LibreOffice n'est pas disponible ou échoue.
    """
    if not _SOFFICE.exists():
        return None
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_docx = Path(tmpdir) / "fiche.docx"
            tmp_docx.write_bytes(docx_bytes)
            subprocess.run(
                [str(_SOFFICE), "--headless", "--convert-to", "odt", "--outdir", tmpdir, str(tmp_docx)],
                check=True,
                capture_output=True,
                timeout=30,
            )
            tmp_odt = Path(tmpdir) / "fiche.odt"
            if tmp_odt.exists():
                return tmp_odt.read_bytes()
    except Exception:
        pass
    return None


# ---------------------------------------------------------------------------
# Écriture sur disque
# ---------------------------------------------------------------------------

def save_normal(
    patient: dict,
    code: str,
    out_dir: Path,
    password: str,
) -> Path:
    """
    Génère le .docx complet, le chiffre (AES-256) et l'écrit dans out_dir.
    Génère aussi le .odt chiffré si LibreOffice est disponible.
    Nom fichier : <Prenom_Nom>_<code>.docx / .odt
    Retourne le chemin du .docx créé.
    """
    identite = patient.get("identite") or {}
    raw_name = identite.get("prenom_nom") or "Patient"
    safe = _safe_name(raw_name)
    docx_bytes = build_docx(patient)

    path = out_dir / f"{safe}_{code}.docx"
    path.write_bytes(encrypt_bytes(docx_bytes, password))

    # ODT chiffré (en parallèle du DOCX)
    odt_bytes = _docx_to_odt_bytes(docx_bytes)
    if odt_bytes:
        odt_path = out_dir / f"{safe}_{code}.odt"
        odt_path.write_bytes(encrypt_bytes(odt_bytes, password))

    return path


def save_anonymized(
    patient_anon: dict,
    code: str,
    out_dir: Path,
) -> Path:
    """
    Génère le .docx pseudonymisé et l'écrit en clair dans out_dir.
    Génère aussi le .odt en clair si LibreOffice est disponible.
    Nom fichier : <code>.docx / .odt
    Retourne le chemin du .docx créé.
    """
    docx_bytes = build_docx(patient_anon)

    path = out_dir / f"{code}.docx"
    path.write_bytes(docx_bytes)

    # ODT en clair (pseudonymisé)
    odt_bytes = _docx_to_odt_bytes(docx_bytes)
    if odt_bytes:
        odt_path = out_dir / f"{code}.odt"
        odt_path.write_bytes(odt_bytes)

    return path
