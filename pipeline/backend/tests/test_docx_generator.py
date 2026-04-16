"""
Tests unitaires pour backend/docx_generator.py
Run: pytest pipeline/backend/tests/test_docx_generator.py -v
"""
import io
import sys
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[3]))

from pipeline.backend.crypto import decrypt_bytes
from pipeline.backend.docx_generator import (
    _safe_name,
    _v,
    build_docx,
    save_anonymized,
    save_normal,
)
from pipeline.backend.pseudonymizer import anonymize


# ---------------------------------------------------------------------------
# Fixture patient complet
# ---------------------------------------------------------------------------

FULL_PATIENT = {
    "date_premiere_seance": "2024-01-15",
    "identite": {
        "prenom_nom": "Marie Dupont",
        "date_naissance": "1980-05-12",
        "profession": "Enseignante",
        "enfants": "2",
        "telephone": "079 111 22 33",
        "mail": "marie@example.com",
        "adresse": "Rue de Berne 1, 1000 Lausanne",
        "canal_communication": "SMS",
    },
    "motif_consultation": "Douleurs lombaires",
    "symptome": {
        "manifestation": "Douleur aiguë",
        "depuis": "3 mois",
        "aggravation_amelioration": "Aggravé par la station assise",
    },
    "histoire_personnelle": "RAS",
    "antecedents_personnels_familiaux": "Hypertension (père)",
    "traitements_en_cours": None,
    "stress_fatigue_psychique": "Modéré",
    "systeme_digestif": "Normal",
    "sommeil": "Perturbé",
    "cardio_vasculaire": {"observations": "Normal", "anticoagulants": None},
    "cycle_menstruel": "Régulier",
    "pouls_langue": "Pouls tendu, langue pâle",
    "seances": [
        {"numero": 1, "date": "2024-01-15", "bilan_ttt": "Première séance, acupuncture lombaire"},
        {"numero": 2, "date": "2024-01-29", "bilan_ttt": "Amélioration 30%"},
        {"numero": 3, "date": None, "bilan_ttt": None},
        {"numero": 4, "date": None, "bilan_ttt": None},
        {"numero": 5, "date": None, "bilan_ttt": None},
        {"numero": 6, "date": None, "bilan_ttt": None},
    ],
}

NULL_PATIENT = {
    "date_premiere_seance": None,
    "identite": {k: None for k in [
        "prenom_nom", "date_naissance", "profession", "enfants",
        "telephone", "mail", "adresse", "canal_communication",
    ]},
    "motif_consultation": None,
    "symptome": {"manifestation": None, "depuis": None, "aggravation_amelioration": None},
    "histoire_personnelle": None,
    "antecedents_personnels_familiaux": None,
    "traitements_en_cours": None,
    "stress_fatigue_psychique": None,
    "systeme_digestif": None,
    "sommeil": None,
    "cardio_vasculaire": {"observations": None, "anticoagulants": None},
    "cycle_menstruel": None,
    "pouls_langue": None,
    "seances": [{"numero": i, "date": None, "bilan_ttt": None} for i in range(1, 7)],
}


def open_docx(data: bytes) -> Document:
    return Document(io.BytesIO(data))


def all_text(doc: Document) -> str:
    """Concatène tout le texte du document (paragraphes + tableaux)."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# _safe_name
# ---------------------------------------------------------------------------

class TestSafeName:
    def test_spaces_to_underscores(self):
        assert _safe_name("Marie Dupont") == "Marie_Dupont"

    def test_multiple_spaces(self):
        assert _safe_name("Jean  Pierre Martin") == "Jean_Pierre_Martin"

    def test_hyphen_preserved(self):
        assert "Anne-Sophie" in _safe_name("Anne-Sophie Roux")

    def test_leading_trailing_stripped(self):
        assert _safe_name("  Alice  ") == "Alice"


# ---------------------------------------------------------------------------
# _v (valeur ou dash)
# ---------------------------------------------------------------------------

class TestV:
    def test_none_returns_dash(self):
        assert _v(None) == "—"

    def test_string_returned_as_is(self):
        assert _v("Bonjour") == "Bonjour"

    def test_integer_stringified(self):
        assert _v(42) == "42"

    def test_empty_string_returned_as_is(self):
        assert _v("") == ""


# ---------------------------------------------------------------------------
# build_docx — format
# ---------------------------------------------------------------------------

class TestBuildDocxFormat:
    def test_returns_bytes(self):
        result = build_docx(FULL_PATIENT)
        assert isinstance(result, bytes)

    def test_valid_docx_magic_bytes(self):
        result = build_docx(FULL_PATIENT)
        # DOCX = ZIP → commence par PK\x03\x04
        assert result[:4] == b"PK\x03\x04"

    def test_openable_by_python_docx(self):
        result = build_docx(FULL_PATIENT)
        doc = open_docx(result)
        assert doc is not None

    def test_null_patient_still_valid_docx(self):
        result = build_docx(NULL_PATIENT)
        doc = open_docx(result)
        assert doc is not None


# ---------------------------------------------------------------------------
# build_docx — contenu patient complet
# ---------------------------------------------------------------------------

class TestBuildDocxContent:
    def setup_method(self):
        self.doc = open_docx(build_docx(FULL_PATIENT))
        self.text = all_text(self.doc)

    def test_title_contains_patient_name(self):
        assert "Marie Dupont" in self.text

    def test_date_premiere_seance(self):
        assert "2024-01-15" in self.text

    # Identité
    def test_date_naissance(self):
        assert "1980-05-12" in self.text

    def test_profession(self):
        assert "Enseignante" in self.text

    def test_telephone(self):
        assert "079 111 22 33" in self.text

    def test_mail(self):
        assert "marie@example.com" in self.text

    def test_adresse(self):
        assert "Rue de Berne 1" in self.text

    # Sections cliniques
    def test_motif_consultation(self):
        assert "Douleurs lombaires" in self.text

    def test_symptome_manifestation(self):
        assert "Douleur aiguë" in self.text

    def test_symptome_depuis(self):
        assert "3 mois" in self.text

    def test_histoire_personnelle(self):
        assert "RAS" in self.text

    def test_antecedents(self):
        assert "Hypertension" in self.text

    def test_sommeil(self):
        assert "Perturbé" in self.text

    def test_cardio_observations(self):
        assert "Normal" in self.text

    def test_cycle_menstruel(self):
        assert "Régulier" in self.text

    def test_pouls_langue(self):
        assert "Pouls tendu" in self.text

    # Séances
    def test_seance_date_present(self):
        assert "2024-01-29" in self.text

    def test_seance_bilan_present(self):
        assert "Amélioration 30%" in self.text

    def test_seances_table_exists(self):
        assert len(self.doc.tables) >= 1

    def test_seances_table_has_header_row(self):
        table = self.doc.tables[0]
        header_texts = [cell.text for cell in table.rows[0].cells]
        assert "N°" in header_texts
        assert "Date" in header_texts
        assert "Bilan / Traitement" in header_texts

    def test_seances_table_has_6_data_rows(self):
        table = self.doc.tables[0]
        assert len(table.rows) == 7  # 1 header + 6 séances

    # Null fields → dash
    def test_null_field_shows_dash(self):
        assert "—" in self.text  # traitements_en_cours = None


# ---------------------------------------------------------------------------
# build_docx — patient totalement null
# ---------------------------------------------------------------------------

class TestBuildDocxNullPatient:
    def test_all_nulls_produce_dashes(self):
        doc = open_docx(build_docx(NULL_PATIENT))
        text = all_text(doc)
        assert "—" in text

    def test_title_fallback(self):
        doc = open_docx(build_docx(NULL_PATIENT))
        text = all_text(doc)
        assert "Patient" in text


# ---------------------------------------------------------------------------
# build_docx — version anonymisée
# ---------------------------------------------------------------------------

class TestBuildDocxAnonymized:
    def setup_method(self):
        anon = anonymize(FULL_PATIENT, "P0001")
        self.doc = open_docx(build_docx(anon))
        self.text = all_text(self.doc)

    def test_code_appears_in_title(self):
        assert "P0001" in self.text

    def test_real_name_absent(self):
        assert "Marie Dupont" not in self.text

    def test_real_phone_absent(self):
        assert "079 111 22 33" not in self.text

    def test_clinical_data_preserved(self):
        assert "Douleurs lombaires" in self.text
        assert "Amélioration 30%" in self.text


# ---------------------------------------------------------------------------
# save_normal
# ---------------------------------------------------------------------------

class TestSaveNormal:
    def test_creates_file(self, tmp_path):
        path = save_normal(FULL_PATIENT, "P0001", tmp_path, "secret")
        assert path.exists()

    def test_filename_format(self, tmp_path):
        path = save_normal(FULL_PATIENT, "P0001", tmp_path, "secret")
        assert path.name == "Marie_Dupont_P0001.docx"

    def test_file_is_encrypted(self, tmp_path):
        path = save_normal(FULL_PATIENT, "P0001", tmp_path, "secret")
        raw = path.read_bytes()
        # Chiffré → pas de magic bytes DOCX (PK\x03\x04)
        assert raw[:4] != b"PK\x03\x04"
        # Et ne contient pas le nom en clair
        assert b"Marie Dupont" not in raw

    def test_decrypted_is_valid_docx(self, tmp_path):
        path = save_normal(FULL_PATIENT, "P0001", tmp_path, "secret")
        plaintext = decrypt_bytes(path.read_bytes(), "secret")
        doc = open_docx(plaintext)
        assert "Marie Dupont" in all_text(doc)

    def test_wrong_password_cannot_open(self, tmp_path):
        path = save_normal(FULL_PATIENT, "P0001", tmp_path, "correct")
        with pytest.raises(ValueError):
            decrypt_bytes(path.read_bytes(), "wrong")

    def test_returns_path(self, tmp_path):
        result = save_normal(FULL_PATIENT, "P0001", tmp_path, "pw")
        assert isinstance(result, Path)


# ---------------------------------------------------------------------------
# save_anonymized
# ---------------------------------------------------------------------------

class TestSaveAnonymized:
    def test_creates_file(self, tmp_path):
        anon = anonymize(FULL_PATIENT, "P0001")
        path = save_anonymized(anon, "P0001", tmp_path)
        assert path.exists()

    def test_filename_is_code_dot_docx(self, tmp_path):
        anon = anonymize(FULL_PATIENT, "P0001")
        path = save_anonymized(anon, "P0001", tmp_path)
        assert path.name == "P0001.docx"

    def test_file_is_readable_docx(self, tmp_path):
        anon = anonymize(FULL_PATIENT, "P0001")
        path = save_anonymized(anon, "P0001", tmp_path)
        doc = Document(str(path))
        assert doc is not None

    def test_file_is_not_encrypted(self, tmp_path):
        anon = anonymize(FULL_PATIENT, "P0001")
        path = save_anonymized(anon, "P0001", tmp_path)
        # Fichier DOCX non chiffré → commence par PK
        assert path.read_bytes()[:4] == b"PK\x03\x04"

    def test_no_real_identity_in_file(self, tmp_path):
        anon = anonymize(FULL_PATIENT, "P0001")
        path = save_anonymized(anon, "P0001", tmp_path)
        raw = path.read_bytes()
        assert b"marie@example.com" not in raw
        assert b"079 111 22 33" not in raw

    def test_returns_path(self, tmp_path):
        anon = anonymize(FULL_PATIENT, "P0001")
        result = save_anonymized(anon, "P0001", tmp_path)
        assert isinstance(result, Path)
